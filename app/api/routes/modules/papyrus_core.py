"""Papyrus — API routes.

CRUD for documents, revisions, templates, doc types,
distribution lists, arborescence nodes, share links.
Workflow transitions (submit, approve, reject, publish).
Export (PDF, DOCX). Revision diff.
"""

import io
import json
import logging
from typing import Optional
from uuid import UUID

from fastapi import APIRouter, Depends, File, HTTPException, Query, Request, Response, UploadFile
from fastapi.responses import StreamingResponse
from sqlalchemy.ext.asyncio import AsyncSession

# Graceful fallback for optional export dependencies
try:
    from weasyprint import HTML as WeasyHTML
except ImportError:
    WeasyHTML = None

try:
    import docx as python_docx
except ImportError:
    python_docx = None

from app.core.database import get_db
from app.api.deps import get_current_entity, get_current_user, require_module_enabled, require_permission

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/v1/documents", tags=["papyrus"], dependencies=[require_module_enabled("papyrus")])


@router.get(
    "/papyrus/forms",
    dependencies=[require_permission("document.read")],
    summary="List Papyrus forms",
)
async def list_papyrus_forms(
    entity_id: UUID = Depends(get_current_entity),
    db: AsyncSession = Depends(get_db),
):
    from app.services.modules.papyrus_forms_service import list_forms
    return await list_forms(entity_id=entity_id, db=db)


@router.post(
    "/papyrus/forms",
    dependencies=[require_permission("document.edit")],
    summary="Create a Papyrus form",
)
async def create_papyrus_form(
    body: dict,
    entity_id: UUID = Depends(get_current_entity),
    current_user=Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    from app.schemas.papyrus import PapyrusFormCreate
    from app.services.modules.papyrus_forms_service import create_form

    parsed = PapyrusFormCreate(**body)
    return await create_form(entity_id=entity_id, created_by=current_user.id, body=parsed, db=db)


@router.post(
    "/papyrus/forms/import/epicollect",
    dependencies=[require_permission("document.edit")],
    summary="Import an EpiCollect5 project into a Papyrus form",
)
async def import_epicollect_form(
    body: dict,
    entity_id: UUID = Depends(get_current_entity),
    current_user=Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    from app.schemas.papyrus import PapyrusEpiCollectImport
    from app.services.modules.papyrus_forms_service import import_epicollect_form as svc_import

    parsed = PapyrusEpiCollectImport(**body)
    return await svc_import(entity_id=entity_id, created_by=current_user.id, body=parsed, db=db)


@router.get(
    "/papyrus/forms/{form_id}",
    dependencies=[require_permission("document.read")],
    summary="Get a Papyrus form",
)
async def get_papyrus_form(
    form_id: UUID,
    entity_id: UUID = Depends(get_current_entity),
    db: AsyncSession = Depends(get_db),
):
    from app.services.modules.papyrus_forms_service import get_form
    return await get_form(form_id=form_id, entity_id=entity_id, db=db)


@router.get(
    "/papyrus/forms/{form_id}/export/epicollect",
    dependencies=[require_permission("document.read")],
    summary="Export a Papyrus form to EpiCollect5 JSON",
)
async def export_epicollect_form(
    form_id: UUID,
    entity_id: UUID = Depends(get_current_entity),
    db: AsyncSession = Depends(get_db),
):
    from app.services.modules.papyrus_forms_service import export_epicollect_form as svc_export
    return await svc_export(form_id=form_id, entity_id=entity_id, db=db)


@router.patch(
    "/papyrus/forms/{form_id}",
    dependencies=[require_permission("document.edit")],
    summary="Update a Papyrus form",
)
async def update_papyrus_form(
    form_id: UUID,
    body: dict,
    entity_id: UUID = Depends(get_current_entity),
    db: AsyncSession = Depends(get_db),
):
    from app.schemas.papyrus import PapyrusFormUpdate
    from app.services.modules.papyrus_forms_service import update_form

    parsed = PapyrusFormUpdate(**body)
    return await update_form(form_id=form_id, entity_id=entity_id, body=parsed, db=db)


@router.get(
    "/papyrus/forms/{form_id}/submissions",
    dependencies=[require_permission("document.read")],
    summary="List Papyrus external submissions",
)
async def list_papyrus_submissions(
    form_id: UUID,
    entity_id: UUID = Depends(get_current_entity),
    db: AsyncSession = Depends(get_db),
):
    from app.services.modules.papyrus_forms_service import list_submissions
    return await list_submissions(form_id=form_id, entity_id=entity_id, db=db)


@router.post(
    "/papyrus/forms/{form_id}/external-links",
    dependencies=[require_permission("document.edit")],
    summary="Create a Papyrus external submission link",
)
async def create_papyrus_external_link(
    form_id: UUID,
    body: dict,
    entity_id: UUID = Depends(get_current_entity),
    current_user=Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    from app.schemas.papyrus import PapyrusExternalLinkCreate
    from app.services.modules.papyrus_forms_service import create_external_link

    parsed = PapyrusExternalLinkCreate(**body)
    return await create_external_link(
        form_id=form_id,
        entity_id=entity_id,
        created_by=current_user.id,
        body=parsed,
        db=db,
    )


@router.delete(
    "/papyrus/forms/{form_id}/external-links/{token_id}",
    dependencies=[require_permission("document.edit")],
    summary="Revoke a Papyrus external submission link",
)
async def revoke_papyrus_external_link(
    form_id: UUID,
    token_id: str,
    entity_id: UUID = Depends(get_current_entity),
    db: AsyncSession = Depends(get_db),
):
    from app.services.modules.papyrus_forms_service import revoke_external_link
    return await revoke_external_link(form_id=form_id, token_id=token_id, entity_id=entity_id, db=db)


@router.get(
    "/papyrus/ext/forms/{form_id}",
    summary="Consume a Papyrus external form link",
)
async def consume_papyrus_external_form(
    form_id: UUID,
    token: str = Query(...),
    request: Request = None,
    db: AsyncSession = Depends(get_db),
):
    from app.services.modules.papyrus_forms_service import consume_external_form
    request_ip = request.client.host if request and request.client else None
    return await consume_external_form(form_id=form_id, token=token, request_ip=request_ip, db=db)


@router.post(
    "/papyrus/ext/forms/{form_id}/submit",
    summary="Submit a Papyrus external form response",
)
async def submit_papyrus_external_form(
    form_id: UUID,
    body: dict,
    token: str = Query(...),
    request: Request = None,
    db: AsyncSession = Depends(get_db),
):
    from app.schemas.papyrus import PapyrusExternalSubmissionCreate
    from app.services.modules.papyrus_forms_service import submit_external_form

    parsed = PapyrusExternalSubmissionCreate(**body)
    request_ip = request.client.host if request and request.client else None
    return await submit_external_form(form_id=form_id, token=token, request_ip=request_ip, body=parsed, db=db)


# ═══════════════════════════════════════════════════════════════════════════════
# Documents CRUD
# ═══════════════════════════════════════════════════════════════════════════════


@router.get(
    "/",
    dependencies=[require_permission("document.read")],
    summary="List documents",
)
async def list_documents(
    project_id: Optional[str] = None,
    doc_type_id: Optional[str] = None,
    status: Optional[str] = None,
    classification: Optional[str] = None,
    arborescence_node_id: Optional[str] = None,
    search: Optional[str] = None,
    page: int = Query(1, ge=1),
    page_size: int = Query(25, ge=1, le=100),
    entity_id: UUID = Depends(get_current_entity),
    current_user=Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    from app.services.modules.papyrus_document_service import list_documents as svc_list
    return await svc_list(
        entity_id=entity_id,
        bu_id=getattr(current_user, "bu_id", None),
        project_id=project_id,
        doc_type_id=doc_type_id,
        status=status,
        classification=classification,
        arborescence_node_id=arborescence_node_id,
        search=search,
        page=page,
        page_size=page_size,
        db=db,
    )


@router.post(
    "/",
    dependencies=[require_permission("document.create")],
    summary="Create a new document",
)
async def create_document(
    body: dict,  # Will use schema once created by agent
    entity_id: UUID = Depends(get_current_entity),
    current_user=Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    from app.schemas.papyrus_document import DocumentCreate
    from app.services.modules.papyrus_document_service import create_document as svc_create

    parsed = DocumentCreate(**body)
    return await svc_create(
        body=parsed,
        entity_id=entity_id,
        bu_id=getattr(current_user, "bu_id", None),
        created_by=current_user.id,
        db=db,
    )


# IMPORTANT: literal routes MUST be before /{doc_id} to avoid UUID match conflict


@router.get(
    "/share/{token}",
    summary="Consume a share link (public, no auth required)",
)
async def consume_share_link(
    token: str,
    db: AsyncSession = Depends(get_db),
):
    """Public endpoint for external access via share link.

    Validates the token, checks expiry and access limits,
    increments access_count, and returns document metadata + revision content.
    Returns 401 if OTP is required.
    """
    from app.services.modules.papyrus_document_service import consume_share_link as svc_consume
    return await svc_consume(token=token, db=db)


@router.get(
    "/counts",
    dependencies=[require_permission("document.read")],
    summary="Get document status counts",
)
async def get_document_counts(
    entity_id: UUID = Depends(get_current_entity),
    db: AsyncSession = Depends(get_db),
):
    from app.services.modules.papyrus_document_service import get_document_counts as svc_counts
    return await svc_counts(entity_id=entity_id, db=db)


@router.get("/templates", dependencies=[require_permission("document.read")], summary="List templates")
async def list_templates_early(
    doc_type_id: Optional[str] = None, entity_id: UUID = Depends(get_current_entity), db: AsyncSession = Depends(get_db),
):
    from app.services.modules.papyrus_document_service import list_templates as svc_list
    return await svc_list(entity_id=entity_id, doc_type_id=UUID(doc_type_id) if doc_type_id else None, db=db)

@router.post("/templates", dependencies=[require_permission("template.create")], summary="Create a template")
async def create_template_early(
    body: dict, entity_id: UUID = Depends(get_current_entity), current_user=Depends(get_current_user), db: AsyncSession = Depends(get_db),
):
    from app.schemas.papyrus_document import TemplateCreate
    from app.services.modules.papyrus_document_service import create_template as svc_create
    parsed = TemplateCreate(**body)
    return await svc_create(body=parsed, entity_id=entity_id, created_by=current_user.id, db=db)


@router.get(
    "/types",
    dependencies=[require_permission("document.read")],
    summary="List document types",
)
async def list_doc_types(
    entity_id: UUID = Depends(get_current_entity),
    db: AsyncSession = Depends(get_db),
):
    from app.services.modules.papyrus_document_service import list_doc_types as svc_list
    return await svc_list(entity_id=entity_id, db=db)


@router.post(
    "/types",
    dependencies=[require_permission("document.admin")],
    summary="Create a document type",
)
async def create_doc_type(
    body: dict,
    entity_id: UUID = Depends(get_current_entity),
    current_user=Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    from app.schemas.papyrus_document import DocTypeCreate
    from app.services.modules.papyrus_document_service import create_doc_type as svc_create
    from fastapi import HTTPException
    from pydantic import ValidationError

    try:
        parsed = DocTypeCreate(**body)
    except ValidationError as exc:
        raise HTTPException(status_code=422, detail=exc.errors()) from exc
    return await svc_create(body=parsed, entity_id=entity_id, created_by=current_user.id, db=db)


@router.patch(
    "/types/{type_id}",
    dependencies=[require_permission("document.admin")],
    summary="Update a document type",
)
async def update_doc_type(
    type_id: str,
    body: dict,
    entity_id: UUID = Depends(get_current_entity),
    db: AsyncSession = Depends(get_db),
):
    from app.schemas.papyrus_document import DocTypeUpdate
    from app.services.modules.papyrus_document_service import update_doc_type as svc_update

    parsed = DocTypeUpdate(**body)
    return await svc_update(type_id=type_id, body=parsed, entity_id=entity_id, db=db)


@router.delete(
    "/types/{type_id}",
    dependencies=[require_permission("document.admin")],
    summary="Soft-delete a document type (only if no documents reference it)",
)
async def delete_doc_type(
    type_id: str,
    entity_id: UUID = Depends(get_current_entity),
    db: AsyncSession = Depends(get_db),
):
    from app.services.modules.papyrus_document_service import delete_doc_type as svc_delete
    return await svc_delete(type_id=type_id, entity_id=entity_id, db=db)


@router.post(
    "/types/mdr/import",
    dependencies=[require_permission("document.admin")],
    summary="Import Master Document Register (CSV/XLSX)",
)
async def import_mdr(
    file: UploadFile = File(...),
    project_id: Optional[str] = Query(None),
    entity_id: UUID = Depends(get_current_entity),
    current_user=Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Import a Master Document Register file (CSV or XLSX).

    Creates or updates DocType records and optionally creates Document
    placeholders when document_number column is present.
    """
    from app.services.modules.papyrus_document_service import import_mdr as svc_import

    return await svc_import(
        file=file,
        entity_id=entity_id,
        project_id=UUID(project_id) if project_id else None,
        created_by=current_user.id,
        db=db,
    )


@router.get(
    "/{doc_id}",
    dependencies=[require_permission("document.read")],
    summary="Get a document by ID",
)
async def get_document(
    doc_id: UUID,
    entity_id: UUID = Depends(get_current_entity),
    db: AsyncSession = Depends(get_db),
):
    from app.services.modules.papyrus_document_service import get_document as svc_get
    return await svc_get(doc_id, entity_id, db)


@router.patch(
    "/{doc_id}",
    dependencies=[require_permission("document.edit")],
    summary="Update document metadata",
)
async def update_document(
    doc_id: UUID,
    body: dict,
    entity_id: UUID = Depends(get_current_entity),
    db: AsyncSession = Depends(get_db),
):
    from app.schemas.papyrus_document import DocumentUpdate
    from app.services.modules.papyrus_document_service import update_document as svc_update

    parsed = DocumentUpdate(**body)
    return await svc_update(doc_id=doc_id, body=parsed, entity_id=entity_id, db=db)


# ═══════════════════════════════════════════════════════════════════════════════
# Draft saving (autosave)
# ═══════════════════════════════════════════════════════════════════════════════


@router.patch(
    "/{doc_id}/draft",
    dependencies=[require_permission("document.edit")],
    summary="Save draft content (autosave)",
)
async def save_draft(
    doc_id: UUID,
    body: dict,
    entity_id: UUID = Depends(get_current_entity),
    current_user=Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    from app.services.modules.papyrus_document_service import save_draft as svc_save

    return await svc_save(
        doc_id=doc_id,
        content=body.get("content", {}),
        form_data=body.get("form_data", {}),
        yjs_state=body.get("yjs_state"),
        entity_id=entity_id,
        user_id=current_user.id,
        db=db,
    )


# ═══════════════════════════════════════════════════════════════════════════════
# Revisions
# ═══════════════════════════════════════════════════════════════════════════════


@router.get(
    "/{doc_id}/revisions",
    dependencies=[require_permission("document.read")],
    summary="List revisions for a document",
)
async def list_revisions(
    doc_id: UUID,
    entity_id: UUID = Depends(get_current_entity),
    db: AsyncSession = Depends(get_db),
):
    from app.services.modules.papyrus_document_service import list_revisions as svc_list
    return await svc_list(doc_id=doc_id, entity_id=entity_id, db=db)


@router.get(
    "/{doc_id}/revisions/{revision_id}",
    dependencies=[require_permission("document.read")],
    summary="Get a specific revision",
)
async def get_revision(
    doc_id: UUID,
    revision_id: str,
    entity_id: UUID = Depends(get_current_entity),
    db: AsyncSession = Depends(get_db),
):
    from app.services.modules.papyrus_document_service import get_revision as svc_get
    return await svc_get(revision_id, entity_id, db)


@router.post(
    "/{doc_id}/revisions",
    dependencies=[require_permission("document.edit")],
    summary="Create a new revision (advance rev code)",
)
async def create_new_revision(
    doc_id: UUID,
    entity_id: UUID = Depends(get_current_entity),
    current_user=Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    from app.services.modules.papyrus_document_service import create_new_revision as svc_create
    return await svc_create(
        doc_id=doc_id,
        entity_id=entity_id,
        user_id=current_user.id,
        db=db,
    )


@router.get(
    "/{doc_id}/diff",
    dependencies=[require_permission("document.read")],
    summary="Compare two revisions",
)
async def diff_revisions(
    doc_id: UUID,
    rev_a: str = Query(..., description="Revision A ID"),
    rev_b: str = Query(..., description="Revision B ID"),
    entity_id: UUID = Depends(get_current_entity),
    db: AsyncSession = Depends(get_db),
):
    from app.services.modules.papyrus_document_service import diff_revisions as svc_diff
    return await svc_diff(
        doc_id=doc_id,
        rev_a_id=rev_a,
        rev_b_id=rev_b,
        entity_id=entity_id,
        db=db,
    )


@router.get(
    "/{doc_id}/papyrus",
    dependencies=[require_permission("document.read")],
    summary="Get the canonical Papyrus document",
)
async def get_papyrus_document(
    doc_id: UUID,
    version: int | None = Query(None, ge=1),
    entity_id: UUID = Depends(get_current_entity),
    db: AsyncSession = Depends(get_db),
):
    from app.services.modules.papyrus_document_service import get_papyrus_document as svc_get
    return await svc_get(doc_id=doc_id, entity_id=entity_id, db=db, version=version)


@router.get(
    "/{doc_id}/papyrus/versions",
    dependencies=[require_permission("document.read")],
    summary="List Papyrus technical versions for a document",
)
async def list_papyrus_versions(
    doc_id: UUID,
    entity_id: UUID = Depends(get_current_entity),
    db: AsyncSession = Depends(get_db),
):
    from app.services.modules.papyrus_document_service import list_papyrus_versions as svc_list
    return await svc_list(doc_id=doc_id, entity_id=entity_id, db=db)


@router.get(
    "/{doc_id}/papyrus/render",
    dependencies=[require_permission("document.read")],
    summary="Get a rendered Papyrus document with refs and formulas resolved",
)
async def get_rendered_papyrus_document(
    doc_id: UUID,
    version: int | None = Query(None, ge=1),
    entity_id: UUID = Depends(get_current_entity),
    db: AsyncSession = Depends(get_db),
):
    from app.services.modules.papyrus_document_service import get_rendered_papyrus_document as svc_get
    return await svc_get(doc_id=doc_id, entity_id=entity_id, db=db, version=version)


@router.get(
    "/{doc_id}/papyrus/schedule",
    dependencies=[require_permission("document.read")],
    summary="Get Papyrus automated dispatch schedule",
)
async def get_papyrus_schedule(
    doc_id: UUID,
    entity_id: UUID = Depends(get_current_entity),
    db: AsyncSession = Depends(get_db),
):
    from app.services.modules.papyrus_dispatch_service import get_document_schedule as svc_get

    return await svc_get(doc_id=doc_id, entity_id=entity_id, db=db)


@router.put(
    "/{doc_id}/papyrus/schedule",
    dependencies=[require_permission("document.edit")],
    summary="Update Papyrus automated dispatch schedule",
)
async def update_papyrus_schedule(
    doc_id: UUID,
    body: dict,
    entity_id: UUID = Depends(get_current_entity),
    current_user=Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    from app.schemas.papyrus import PapyrusScheduleUpdate
    from app.services.modules.papyrus_dispatch_service import update_document_schedule as svc_update

    parsed = PapyrusScheduleUpdate(**body)
    return await svc_update(
        doc_id=doc_id,
        entity_id=entity_id,
        actor_id=current_user.id,
        body=parsed,
        db=db,
    )


@router.get(
    "/{doc_id}/papyrus/dispatch-runs",
    dependencies=[require_permission("document.read")],
    summary="List Papyrus dispatch runs for a document",
)
async def list_papyrus_dispatch_runs(
    doc_id: UUID,
    limit: int = Query(20, ge=1, le=100),
    entity_id: UUID = Depends(get_current_entity),
    db: AsyncSession = Depends(get_db),
):
    from app.services.modules.papyrus_dispatch_service import list_dispatch_runs as svc_list

    return await svc_list(doc_id=doc_id, entity_id=entity_id, db=db, limit=limit)


@router.post(
    "/{doc_id}/papyrus/dispatch-run-now",
    dependencies=[require_permission("document.edit")],
    summary="Trigger Papyrus dispatch immediately",
)
async def run_papyrus_dispatch_now(
    doc_id: UUID,
    entity_id: UUID = Depends(get_current_entity),
    current_user=Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    from app.services.modules.papyrus_dispatch_service import dispatch_document_now as svc_run

    return await svc_run(
        doc_id=doc_id,
        entity_id=entity_id,
        triggered_by=current_user.id,
        db=db,
    )


# ═══════════════════════════════════════════════════════════════════════════════
# Dynamic Workflow (FSM-driven)
# ═══════════════════════════════════════════════════════════════════════════════


@router.get(
    "/{doc_id}/workflow-state",
    dependencies=[require_permission("document.read")],
    summary="Get workflow state, available transitions, and history",
)
async def get_workflow_state(
    doc_id: UUID,
    entity_id: UUID = Depends(get_current_entity),
    current_user=Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Return the dynamic workflow state for a document.

    Returns current state, available transitions (with labels,
    required roles, comment requirements), and transition history.
    """
    from app.services.modules.papyrus_document_service import get_workflow_state as svc_get

    return await svc_get(
        doc_id=doc_id,
        entity_id=entity_id,
        user_id=current_user.id,
        db=db,
    )


@router.post(
    "/{doc_id}/transition",
    dependencies=[require_permission("document.edit")],
    summary="Execute a workflow transition",
)
async def execute_transition(
    doc_id: UUID,
    body: dict,
    entity_id: UUID = Depends(get_current_entity),
    current_user=Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Execute a workflow transition with optional comment.

    Body: {"to_state": "approved", "comment": "Looks good"}
    Returns the updated workflow state.
    """
    from app.services.modules.papyrus_document_service import execute_transition as svc_transition

    to_state = body.get("to_state")
    if not to_state:
        raise HTTPException(400, "to_state is required")

    return await svc_transition(
        doc_id=doc_id,
        to_state=to_state,
        comment=body.get("comment"),
        actor_id=current_user.id,
        entity_id=entity_id,
        db=db,
    )


# ═══════════════════════════════════════════════════════════════════════════════
# Workflow transitions (legacy — kept for backward compatibility)
# ═══════════════════════════════════════════════════════════════════════════════


@router.post(
    "/{doc_id}/submit",
    dependencies=[require_permission("document.submit")],
    summary="Submit document for validation",
)
async def submit_document(
    doc_id: UUID,
    body: dict | None = None,
    entity_id: UUID = Depends(get_current_entity),
    current_user=Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    from app.services.modules.papyrus_document_service import submit_document as svc_submit
    return await svc_submit(
        doc_id=doc_id,
        comment=(body or {}).get("comment"),
        entity_id=entity_id,
        actor_id=current_user.id,
        db=db,
    )


@router.post(
    "/{doc_id}/approve",
    dependencies=[require_permission("document.approve")],
    summary="Approve document",
)
async def approve_document(
    doc_id: UUID,
    body: dict | None = None,
    entity_id: UUID = Depends(get_current_entity),
    current_user=Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    from app.services.modules.papyrus_document_service import approve_document as svc_approve
    return await svc_approve(
        doc_id=doc_id,
        comment=(body or {}).get("comment"),
        entity_id=entity_id,
        actor_id=current_user.id,
        db=db,
    )


@router.post(
    "/{doc_id}/reject",
    dependencies=[require_permission("document.reject")],
    summary="Reject document",
)
async def reject_document(
    doc_id: UUID,
    body: dict,
    entity_id: UUID = Depends(get_current_entity),
    current_user=Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    from app.services.modules.papyrus_document_service import reject_document as svc_reject
    reason = body.get("reason", "")
    if not reason:
        raise HTTPException(400, "Rejection reason is required")
    return await svc_reject(
        doc_id=doc_id,
        reason=reason,
        entity_id=entity_id,
        actor_id=current_user.id,
        db=db,
    )


@router.post(
    "/{doc_id}/publish",
    dependencies=[require_permission("document.publish")],
    summary="Publish document (D-083: manual after approval)",
)
async def publish_document(
    doc_id: UUID,
    body: dict | None = None,
    entity_id: UUID = Depends(get_current_entity),
    current_user=Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    from app.services.modules.papyrus_document_service import publish_document as svc_publish
    return await svc_publish(
        doc_id=doc_id,
        distribution_list_ids=[
            UUID(dl_id) for dl_id in (body or {}).get("distribution_list_ids", [])
        ],
        entity_id=entity_id,
        actor_id=current_user.id,
        db=db,
    )


@router.post(
    "/{doc_id}/obsolete",
    dependencies=[require_permission("document.publish")],
    summary="Obsolete a published document",
)
async def obsolete_document(
    doc_id: UUID,
    body: dict | None = None,
    entity_id: UUID = Depends(get_current_entity),
    current_user=Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    from app.services.modules.papyrus_document_service import obsolete_document as svc_obsolete
    return await svc_obsolete(
        doc_id=doc_id,
        superseded_by=UUID((body or {})["superseded_by"]) if (body or {}).get("superseded_by") else None,
        entity_id=entity_id,
        actor_id=current_user.id,
        db=db,
    )


# ═══════════════════════════════════════════════════════════════════════════════
# Export (PDF / DOCX)
# ═══════════════════════════════════════════════════════════════════════════════


@router.get(
    "/{doc_id}/export/pdf",
    dependencies=[require_permission("document.read")],
    summary="Export document as PDF",
)
async def export_pdf(
    doc_id: UUID,
    revision_id: Optional[str] = None,
    entity_id: UUID = Depends(get_current_entity),
    db: AsyncSession = Depends(get_db),
):
    """Export document as PDF via the centralized PDF template engine."""

    from app.services.modules.papyrus_document_service import get_document, get_revision
    from app.core.pdf_templates import render_pdf
    from app.models.common import Entity, User
    from app.models.papyrus_document import DocType

    doc = await get_document(doc_id, entity_id, db)

    # Resolve revision: use specified revision_id or current_revision_id
    rev_id = revision_id or (str(doc.current_revision_id) if doc.current_revision_id else None)
    revision = None
    if rev_id:
        try:
            revision = await get_revision(rev_id, entity_id, db)
        except Exception:
            logger.warning("Revision %s not found, exporting with empty content", rev_id)

    # Build HTML from revision content (BlockNote JSON → simple HTML)
    content_json = revision.content if revision else {}
    form_data = revision.form_data if revision else {}
    html_body = _render_content_to_html(content_json)
    form_html = _render_form_data_to_html(form_data)

    entity = await db.get(Entity, entity_id)
    author = await db.get(User, doc.created_by) if doc.created_by else None
    doc_type = await db.get(DocType, doc.doc_type_id) if doc.doc_type_id else None
    variables = {
        "document_number": doc.number,
        "document_title": doc.title,
        "document_body": f"{form_html}{html_body}",
        "author_name": f"{author.first_name} {author.last_name}".strip() if author else "--",
        "revision": revision.rev_code if revision else "-",
        "status": doc.status,
        "entity": {"name": entity.name if entity else ""},
        "generated_at": datetime.now(timezone.utc).strftime("%d/%m/%Y %H:%M"),
        "document_language": getattr(doc, "language", "fr"),
        "classification": getattr(doc, "classification", ""),
        "doc_type_name": doc_type.name if doc_type else "",
    }

    try:
        pdf_bytes = await render_pdf(
            db,
            slug="document.export",
            entity_id=entity_id,
            language=getattr(doc, "language", "fr") or "fr",
            variables=variables,
        )
    except RuntimeError as exc:
        raise HTTPException(status_code=503, detail=str(exc)) from exc

    if not pdf_bytes:
        raise HTTPException(
            status_code=404,
            detail="Template PDF 'document.export' introuvable. Creez-le dans Parametres > Modeles PDF.",
        )

    return Response(
        content=pdf_bytes,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{doc.number}.pdf"'},
    )


@router.get(
    "/{doc_id}/export/docx",
    dependencies=[require_permission("document.read")],
    summary="Export document as Word (.docx)",
)
async def export_docx(
    doc_id: UUID,
    revision_id: Optional[str] = None,
    entity_id: UUID = Depends(get_current_entity),
    db: AsyncSession = Depends(get_db),
):
    """Export document as DOCX using python-docx.

    Creates a Word document with title, metadata, form data table,
    and content paragraphs from the current (or specified) revision.
    """
    if python_docx is None:
        raise HTTPException(
            status_code=501,
            detail="DOCX export not available — install python-docx",
        )

    from app.services.modules.papyrus_document_service import get_document, get_revision

    doc = await get_document(doc_id, entity_id, db)

    # Resolve revision
    rev_id = revision_id or (str(doc.current_revision_id) if doc.current_revision_id else None)
    revision = None
    if rev_id:
        try:
            revision = await get_revision(rev_id, entity_id, db)
        except Exception:
            logger.warning("Revision %s not found, exporting with empty content", rev_id)

    try:
        word_doc = python_docx.Document()

        # Title
        word_doc.add_heading(doc.title, level=0)

        # Metadata paragraph
        meta_para = word_doc.add_paragraph()
        meta_para.add_run(f"Document: {doc.number}").bold = True
        meta_para.add_run(f"    Rev: {revision.rev_code if revision else '-'}")
        meta_para.add_run(f"    Status: {doc.status}")

        word_doc.add_paragraph("")  # spacer

        # Form data as table (if present)
        form_data = revision.form_data if revision else {}
        if form_data and isinstance(form_data, dict):
            word_doc.add_heading("Form Data", level=1)
            table = word_doc.add_table(rows=1, cols=2)
            table.style = "Table Grid"
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "Field"
            hdr_cells[1].text = "Value"
            for key, value in form_data.items():
                row_cells = table.add_row().cells
                row_cells[0].text = str(key)
                row_cells[1].text = str(value) if value is not None else ""
            word_doc.add_paragraph("")  # spacer

        # Content from revision (BlockNote JSON → paragraphs)
        content_json = revision.content if revision else {}
        _add_content_to_docx(word_doc, content_json)

        # Write to buffer
        buffer = io.BytesIO()
        word_doc.save(buffer)
        buffer.seek(0)

    except Exception as exc:
        logger.exception("DOCX generation failed for doc %s", doc_id)
        raise HTTPException(500, f"DOCX generation failed: {exc}") from exc

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{doc.number}.docx"'},
    )


# ═══════════════════════════════════════════════════════════════════════════════
# Doc Types
# ═══════════════════════════════════════════════════════════════════════════════


@router.patch(
    "/templates/{template_id}",
    dependencies=[require_permission("template.edit")],
    summary="Update a template",
)
async def update_template(
    template_id: str,
    body: dict,
    entity_id: UUID = Depends(get_current_entity),
    db: AsyncSession = Depends(get_db),
):
    from app.schemas.papyrus_document import TemplateUpdate
    from app.services.modules.papyrus_document_service import update_template as svc_update

    parsed = TemplateUpdate(**body)
    return await svc_update(template_id=template_id, body=parsed, entity_id=entity_id, db=db)


# ═══════════════════════════════════════════════════════════════════════════════
# Template Fields CRUD
# ═══════════════════════════════════════════════════════════════════════════════


@router.get(
    "/templates/{template_id}/fields",
    dependencies=[require_permission("document.read")],
    summary="List fields for a template",
)
async def list_template_fields(
    template_id: str,
    entity_id: UUID = Depends(get_current_entity),
    db: AsyncSession = Depends(get_db),
):
    from app.services.modules.papyrus_document_service import list_template_fields as svc_list
    return await svc_list(template_id=template_id, entity_id=entity_id, db=db)


@router.post(
    "/templates/{template_id}/fields",
    dependencies=[require_permission("template.create")],
    summary="Create a template field",
)
async def create_template_field(
    template_id: str,
    body: dict,
    entity_id: UUID = Depends(get_current_entity),
    db: AsyncSession = Depends(get_db),
):
    from app.schemas.papyrus_document import TemplateFieldCreate
    from app.services.modules.papyrus_document_service import create_template_field as svc_create

    parsed = TemplateFieldCreate(**body)
    return await svc_create(template_id=template_id, body=parsed, entity_id=entity_id, db=db)


@router.patch(
    "/templates/{template_id}/fields/{field_id}",
    dependencies=[require_permission("template.edit")],
    summary="Update a template field",
)
async def update_template_field(
    template_id: str,
    field_id: str,
    body: dict,
    entity_id: UUID = Depends(get_current_entity),
    db: AsyncSession = Depends(get_db),
):
    from app.schemas.papyrus_document import TemplateFieldUpdate
    from app.services.modules.papyrus_document_service import update_template_field as svc_update

    parsed = TemplateFieldUpdate(**body)
    return await svc_update(
        template_id=template_id, field_id=field_id, body=parsed, entity_id=entity_id, db=db,
    )


@router.delete(
    "/templates/{template_id}/fields/{field_id}",
    dependencies=[require_permission("template.edit")],
    summary="Delete a template field",
)
async def delete_template_field(
    template_id: str,
    field_id: str,
    entity_id: UUID = Depends(get_current_entity),
    db: AsyncSession = Depends(get_db),
):
    from app.services.modules.papyrus_document_service import delete_template_field as svc_delete
    return await svc_delete(template_id=template_id, field_id=field_id, entity_id=entity_id, db=db)


# ═══════════════════════════════════════════════════════════════════════════════
# Distribution Lists
# ═══════════════════════════════════════════════════════════════════════════════


@router.get(
    "/distribution-lists",
    dependencies=[require_permission("document.admin")],
    summary="List distribution lists",
)
async def list_distribution_lists(
    doc_type_id: Optional[str] = None,
    entity_id: UUID = Depends(get_current_entity),
    db: AsyncSession = Depends(get_db),
):
    from app.services.modules.papyrus_document_service import list_distribution_lists as svc_list
    return await svc_list(
        entity_id=entity_id,
        doc_type_id=UUID(doc_type_id) if doc_type_id else None,
        db=db,
    )


@router.post(
    "/distribution-lists",
    dependencies=[require_permission("document.admin")],
    summary="Create a distribution list",
)
async def create_distribution_list(
    body: dict,
    entity_id: UUID = Depends(get_current_entity),
    current_user=Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    from app.schemas.papyrus_document import DistributionListCreate
    from app.services.modules.papyrus_document_service import create_distribution_list as svc_create

    parsed = DistributionListCreate(**body)
    return await svc_create(body=parsed, entity_id=entity_id, created_by=current_user.id, db=db)


@router.patch(
    "/distribution-lists/{list_id}",
    dependencies=[require_permission("document.admin")],
    summary="Update a distribution list",
)
async def update_distribution_list(
    list_id: str,
    body: dict,
    entity_id: UUID = Depends(get_current_entity),
    db: AsyncSession = Depends(get_db),
):
    from app.schemas.papyrus_document import DistributionListUpdate
    from app.services.modules.papyrus_document_service import update_distribution_list as svc_update

    parsed = DistributionListUpdate(**body)
    return await svc_update(list_id=list_id, body=parsed, entity_id=entity_id, db=db)


@router.delete(
    "/distribution-lists/{list_id}",
    dependencies=[require_permission("document.admin")],
    summary="Soft-delete a distribution list",
)
async def delete_distribution_list(
    list_id: str,
    entity_id: UUID = Depends(get_current_entity),
    db: AsyncSession = Depends(get_db),
):
    from app.services.modules.papyrus_document_service import delete_distribution_list as svc_delete
    return await svc_delete(list_id=list_id, entity_id=entity_id, db=db)


# ═══════════════════════════════════════════════════════════════════════════════
# Arborescence Nodes
# ═══════════════════════════════════════════════════════════════════════════════


@router.get(
    "/arborescence/{project_id}",
    dependencies=[require_permission("document.read")],
    summary="List arborescence nodes for a project",
)
async def list_arborescence_nodes(
    project_id: str,
    entity_id: UUID = Depends(get_current_entity),
    db: AsyncSession = Depends(get_db),
):
    from app.services.modules.papyrus_document_service import list_arborescence_nodes as svc_list
    return await svc_list(project_id=UUID(project_id), entity_id=entity_id, db=db)


@router.post(
    "/arborescence",
    dependencies=[require_permission("document.admin")],
    summary="Create an arborescence node",
)
async def create_arborescence_node(
    body: dict,
    entity_id: UUID = Depends(get_current_entity),
    db: AsyncSession = Depends(get_db),
):
    from app.schemas.papyrus_document import ArborescenceNodeCreate
    from app.services.modules.papyrus_document_service import create_arborescence_node as svc_create

    parsed = ArborescenceNodeCreate(**body)
    return await svc_create(body=parsed, entity_id=entity_id, db=db)


# ═══════════════════════════════════════════════════════════════════════════════
# Share Links
# ═══════════════════════════════════════════════════════════════════════════════


@router.post(
    "/{doc_id}/share",
    dependencies=[require_permission("document.share")],
    summary="Create a temporary share link",
)
async def create_share_link(
    doc_id: UUID,
    body: dict | None = None,
    entity_id: UUID = Depends(get_current_entity),
    current_user=Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    from app.services.modules.papyrus_document_service import create_share_link as svc_create
    body = body or {}
    return await svc_create(
        document_id=UUID(doc_id),
        entity_id=entity_id,
        expires_days=body.get("expires_days", 30),
        otp_required=body.get("otp_required", False),
        max_accesses=body.get("max_accesses"),
        created_by=current_user.id,
        db=db,
    )


# ═══════════════════════════════════════════════════════════════════════════════
# Document Signatures
# ═══════════════════════════════════════════════════════════════════════════════


@router.get(
    "/{doc_id}/signatures",
    dependencies=[require_permission("document.read")],
    summary="List all signatures for a document",
)
async def list_document_signatures(
    doc_id: UUID,
    entity_id: UUID = Depends(get_current_entity),
    db: AsyncSession = Depends(get_db),
):
    from app.services.modules.papyrus_document_service import list_document_signatures as svc_list
    return await svc_list(doc_id=doc_id, entity_id=entity_id, db=db)


# ═══════════════════════════════════════════════════════════════════════════════
# Archive / Delete
# ═══════════════════════════════════════════════════════════════════════════════


@router.post(
    "/{doc_id}/archive",
    dependencies=[require_permission("document.admin")],
    summary="Archive a document (any status → archived)",
)
async def archive_document(
    doc_id: UUID,
    entity_id: UUID = Depends(get_current_entity),
    current_user=Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Archive a document. Admin-only. Document remains consultable via filter."""
    from app.services.modules.papyrus_document_service import archive_document as svc_archive
    return await svc_archive(
        doc_id=doc_id,
        entity_id=entity_id,
        actor_id=current_user.id,
        db=db,
    )


@router.delete(
    "/{doc_id}",
    dependencies=[require_permission("document.delete")],
    summary="Soft-delete a draft document (never submitted)",
)
async def delete_document(
    doc_id: UUID,
    entity_id: UUID = Depends(get_current_entity),
    current_user=Depends(get_current_user),
    db: AsyncSession = Depends(get_db),
):
    """Soft-delete a document. Only allowed for drafts that have never been submitted."""
    from app.services.modules.papyrus_document_service import delete_document as svc_delete
    return await svc_delete(
        doc_id=doc_id,
        entity_id=entity_id,
        actor_id=current_user.id,
        db=db,
    )


# ═══════════════════════════════════════════════════════════════════════════════
# Nomenclature pattern validation
# ═══════════════════════════════════════════════════════════════════════════════


@router.post(
    "/nomenclature/validate",
    dependencies=[require_permission("document.admin")],
    summary="Validate a nomenclature pattern",
)
async def validate_nomenclature_pattern(body: dict):
    from app.services.modules.nomenclature_service import validate_nomenclature_pattern as validate
    pattern = body.get("pattern", "")
    errors = validate(pattern)
    return {"pattern": pattern, "is_valid": len(errors) == 0, "errors": errors}


# ═══════════════════════════════════════════════════════════════════════════════
# Export helpers — BlockNote JSON → HTML / DOCX content
# ═══════════════════════════════════════════════════════════════════════════════


def _extract_text_from_inline(inline_content: list) -> str:
    """Extract plain text from BlockNote inline content array."""
    parts: list[str] = []
    if not isinstance(inline_content, list):
        return str(inline_content) if inline_content else ""
    for item in inline_content:
        if isinstance(item, str):
            parts.append(item)
        elif isinstance(item, dict):
            parts.append(item.get("text", ""))
    return "".join(parts)


def _render_content_to_html(content: dict | list | None) -> str:
    """Convert BlockNote JSON content to simple HTML.

    Supports common block types: paragraph, heading, bulletListItem,
    numberedListItem, table, codeBlock, image.  Unknown types fall back
    to <p> with the extracted text.
    """
    if not content:
        return "<p><em>No content</em></p>"

    blocks: list = []
    if isinstance(content, dict):
        if isinstance(content.get("html"), str) and content.get("html"):
            return content["html"]
        blocks = content.get("blocks", content.get("content", []))
        if not blocks and not isinstance(blocks, list):
            # Might be raw dict with text
            return f"<p>{json.dumps(content, ensure_ascii=False, default=str)[:2000]}</p>"
    elif isinstance(content, list):
        blocks = content

    if not blocks:
        return "<p><em>No content</em></p>"

    html_parts: list[str] = []
    for block in blocks:
        if not isinstance(block, dict):
            html_parts.append(f"<p>{block}</p>")
            continue

        btype = block.get("type", "paragraph")
        if btype == "legacy_payload":
            payload = block.get("payload")
            if isinstance(payload, dict) and isinstance(payload.get("html"), str):
                html_parts.append(payload["html"])
            else:
                html_parts.append(f"<pre>{json.dumps(payload, ensure_ascii=False, default=str)}</pre>")
            continue
        if btype == "formula":
            label = block.get("label") or "Formula"
            html_parts.append(f"<p><strong>{label}:</strong> {block.get('computed_value', '')}</p>")
            continue
        if btype in {"opsflux_kpi", "opsflux_asset", "opsflux_actions", "opsflux_gantt"}:
            label = block.get("label") or btype
            display_value = block.get("display_value")
            resolved = block.get("resolved")
            if display_value is None and resolved is not None:
                display_value = resolved
            html_parts.append(f"<p><strong>{label}:</strong> {json.dumps(display_value, ensure_ascii=False, default=str) if isinstance(display_value, (dict, list)) else display_value}</p>")
            continue
        inline = block.get("content", [])
        text = _extract_text_from_inline(inline)
        props = block.get("props", {})

        if btype == "heading":
            level = props.get("level", 2)
            level = min(max(int(level), 1), 6)
            html_parts.append(f"<h{level}>{text}</h{level}>")
        elif btype == "bulletListItem":
            html_parts.append(f"<li>{text}</li>")
        elif btype == "numberedListItem":
            html_parts.append(f"<li>{text}</li>")
        elif btype == "codeBlock":
            html_parts.append(f"<pre><code>{text}</code></pre>")
        elif btype == "image":
            url = props.get("url", "")
            html_parts.append(f'<img src="{url}" style="max-width:100%;" />')
        elif btype == "table":
            rows = block.get("content", {}).get("rows", [])
            if isinstance(rows, list):
                html_parts.append("<table>")
                for row in rows:
                    html_parts.append("<tr>")
                    cells = row.get("cells", []) if isinstance(row, dict) else []
                    for cell in cells:
                        cell_text = _extract_text_from_inline(cell) if isinstance(cell, list) else str(cell)
                        html_parts.append(f"<td>{cell_text}</td>")
                    html_parts.append("</tr>")
                html_parts.append("</table>")
        else:
            # Default: paragraph
            html_parts.append(f"<p>{text}</p>" if text else "")

        # Recurse into children blocks
        children = block.get("children", [])
        if children:
            html_parts.append(_render_content_to_html(children))

    return "\n".join(html_parts)


def _render_form_data_to_html(form_data: dict | None) -> str:
    """Render form_data dict as an HTML table."""
    if not form_data or not isinstance(form_data, dict):
        return ""
    rows = ""
    for key, value in form_data.items():
        rows += f"<tr><th>{key}</th><td>{value if value is not None else ''}</td></tr>"
    if not rows:
        return ""
    return f"<h2>Form Data</h2><table>{rows}</table>"


def _add_content_to_docx(word_doc, content: dict | list | None) -> None:
    """Add BlockNote JSON content as paragraphs/headings to a python-docx Document."""
    if python_docx is None:
        return

    if not content:
        word_doc.add_paragraph("(No content)")
        return

    blocks: list = []
    if isinstance(content, dict):
        if isinstance(content.get("html"), str) and content.get("html"):
            word_doc.add_paragraph(content["html"])
            return
        blocks = content.get("blocks", content.get("content", []))
    elif isinstance(content, list):
        blocks = content

    if not blocks:
        word_doc.add_paragraph("(No content)")
        return

    for block in blocks:
        if not isinstance(block, dict):
            word_doc.add_paragraph(str(block))
            continue

        btype = block.get("type", "paragraph")
        if btype == "legacy_payload":
            payload = block.get("payload")
            word_doc.add_paragraph(json.dumps(payload, ensure_ascii=False, default=str))
            continue
        if btype == "formula":
            label = block.get("label") or "Formula"
            word_doc.add_paragraph(f"{label}: {block.get('computed_value', '')}")
            continue
        if btype in {"opsflux_kpi", "opsflux_asset", "opsflux_actions", "opsflux_gantt"}:
            label = block.get("label") or btype
            display_value = block.get("display_value")
            resolved = block.get("resolved")
            if display_value is None and resolved is not None:
                display_value = resolved
            word_doc.add_paragraph(
                f"{label}: {json.dumps(display_value, ensure_ascii=False, default=str) if isinstance(display_value, (dict, list)) else display_value}"
            )
            continue
        inline = block.get("content", [])
        text = _extract_text_from_inline(inline)
        props = block.get("props", {})

        if btype == "heading":
            level = props.get("level", 2)
            level = min(max(int(level), 1), 4)  # python-docx supports levels 0-9
            word_doc.add_heading(text, level=level)
        elif btype in ("bulletListItem", "numberedListItem"):
            word_doc.add_paragraph(text, style="List Bullet")
        elif btype == "codeBlock":
            p = word_doc.add_paragraph()
            run = p.add_run(text)
            run.font.name = "Courier New"
            run.font.size = python_docx.shared.Pt(9)
        elif btype == "table":
            rows = block.get("content", {}).get("rows", [])
            if isinstance(rows, list) and rows:
                first_row = rows[0] if isinstance(rows[0], dict) else {}
                n_cols = len(first_row.get("cells", [])) if first_row else 1
                n_cols = max(n_cols, 1)
                table = word_doc.add_table(rows=len(rows), cols=n_cols)
                table.style = "Table Grid"
                for r_idx, row in enumerate(rows):
                    cells = row.get("cells", []) if isinstance(row, dict) else []
                    for c_idx, cell in enumerate(cells):
                        if c_idx < n_cols:
                            cell_text = _extract_text_from_inline(cell) if isinstance(cell, list) else str(cell)
                            table.rows[r_idx].cells[c_idx].text = cell_text
        else:
            if text:
                word_doc.add_paragraph(text)

        # Recurse into children
        children = block.get("children", [])
        if children:
            _add_content_to_docx(word_doc, children)

